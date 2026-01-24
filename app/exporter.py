from __future__ import annotations

import json
from typing import Dict, Any
from pathlib import Path
from docx import Document

SECTION_TITLES = {
    "deviation_summary": "1) суть отклонения",
    "deviation_description_full": "2) описание отклонения (расширенно)",
    "root_causes_deviation": "3) коренные причины отклонения",
    "deviation_cost": "4) стоимость отклонения для бизнеса",
    "deviation_consequences": "5) негативные последствия на бизнес",
    "risk_factors": "6) риск-факторы (по риску)",
    "risk_cost": "7) стоимость риска (p×i)",
    "measures": "8) мероприятия (связь с причинами/факторами)",
    "rsbu_accounting": "9) рсбу: отражение, проводки, проверки",
    "ifrs_accounting": "10) мсфо: отражение, корректировки, проверки",
    "primary_documents": "11) первичные документы и источники",
    "owner_highlights": "12) для собственника/топ-менеджмента",
    "business_questions": "13) вопросы бизнесу (уточнение)",
}

def export_docx(path: str | Path, deviation_id: int, user_input: Dict[str, Any],
                selected: Dict[str, Any], sections: Dict[str, Any],
                chosen_variants: Dict[str, int] | None = None,
                view_mode: Dict[str, str] | None = None) -> Path:
    chosen_variants = chosen_variants or {}
    view_mode = view_mode or {}

    doc = Document()
    doc.add_heading(f"отчёт по отклонению (конструктор) — id {deviation_id}", level=1)

    # meta
    doc.add_paragraph(f"категория отклонения: {selected.get('deviation_category', {}).get('primary_id','')}")
    doc.add_paragraph(f"риск: {selected.get('risk', {}).get('primary_id','')}")

    doc.add_heading("входные данные аудитора", level=2)
    for k, v in user_input.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_page_break()

    # sections
    for key, title in SECTION_TITLES.items():
        if key not in sections:
            continue
        doc.add_heading(title, level=2)

        idx = int(chosen_variants.get(key, 0))
        mode = view_mode.get(key, "full")  # export default = full

        variants = sections[key].get("variants", [])
        if not variants:
            doc.add_paragraph("нет данных")
            continue
        idx = max(0, min(idx, len(variants)-1))
        block = variants[idx].get(mode) or variants[idx].get("full") or variants[idx].get("short") or ""
        for para in str(block).split("\n"):
            doc.add_paragraph(para)

    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out
